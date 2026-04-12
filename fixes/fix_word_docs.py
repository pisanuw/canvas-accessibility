"""
Word document (.docx) accessibility fixes.

Fixes:
  - word.image_alt          : images missing alt text         (AI-assisted)
  - word.image_alt_placeholder : images missing alt text      (placeholder, no AI)
  - word.table_headers      : tables missing header row mark  (fully automatic)
  - word.no_language        : no document language set        (fully automatic)
  - word.heading_order      : heading levels skip             (fully automatic)
  - word.headings_presence  : no heading styles used at all   (fully automatic)
  - word.headings_start_at_one : first heading is not H1      (fully automatic)

Usage (CLI):
  python3 fix_word_docs.py --course-id 1492292 --fix all
  python3 fix_word_docs.py --course-id 1492292 --fix image_alt --file-id 79530490
  python3 fix_word_docs.py --course-id 1492292 --fix headings_presence --dry-run
"""

import argparse
import io
import re
import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))
from fixes.canvas_client import CanvasClient
from fixes.ai_client import describe_image

WORD_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
WORD_EXTS = {".docx", ".doc"}

PLACEHOLDER_ALT = (
    "This image currently does not have a description. "
    "Your instructor will review it and add a description."
)

_FILENAME_RE = re.compile(
    r'\.[a-zA-Z0-9]{2,5}$'
    r'|^[\w\-. ]+\.(png|jpe?g|gif|svg|bmp|webp|tiff?|pdf|docx?|pptx?|xlsx?)$',
    re.IGNORECASE,
)


def _is_filename_alt(text: str) -> bool:
    t = text.strip()
    return bool(t and _FILENAME_RE.search(t))


# ── Fix: image alt text ───────────────────────────────────────────────────────

def fix_image_alt(doc_bytes: bytes, filename: str,
                  placeholder: bool = False) -> tuple[bytes, list[str]]:
    """
    Add alt text to all inline/anchored images that have none.

    placeholder=False (default): call Claude Vision to generate descriptions.
    placeholder=True:            set a standard placeholder string (no AI key needed).

    Returns (updated_docx_bytes, list_of_changes).
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(doc_bytes))
    changes = []

    # Cover both inline shapes (shape._inline.docPr) and anchored/floating
    # shapes (wp:docPr in the XML tree) with a single pass over all docPr nodes.
    for docPr in doc.element.body.iter(qn("wp:docPr")):
        existing = docPr.get("descr", "").strip()
        if existing and not _is_filename_alt(existing):
            continue  # already has a real description

        if placeholder:
            description = PLACEHOLDER_ALT
            docPr.set("descr", description)
            name = docPr.get("name", "image")
            changes.append(f"Set placeholder alt text on image: '{name}'")
        else:
            # Find the parent inline/anchor element to get the image bytes
            try:
                # Walk up to find blip embed rId
                ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
                ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                blip = docPr.getparent().getparent().find(
                    f".//{{{ns_a}}}blip"
                )
                if blip is None:
                    continue
                rId = blip.get(f"{{{ns_r}}}embed")
                if not rId:
                    continue
                # Resolve the part from the document part
                image_part = doc.part.related_parts[rId]
                image_bytes = image_part.blob
                ct = image_part.content_type
            except Exception as e:
                changes.append(f"Could not extract image bytes: {e}")
                continue

            try:
                description = describe_image(image_bytes, ct)
                docPr.set("descr", description)
                docPr.set("title", description[:50])
                label = description[:60] + "…" if len(description) > 60 else description
                changes.append(f"Added alt text: '{label or 'DECORATIVE'}'")
            except Exception as e:
                changes.append(f"Could not describe image: {e}")

    if not changes:
        return doc_bytes, changes

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), changes


# ── Fix: table headers ────────────────────────────────────────────────────────

def fix_table_headers(doc_bytes: bytes) -> tuple[bytes, list[str]]:
    """
    Mark the first row of each table as a header row using tblHeader.
    Returns (updated_docx_bytes, list_of_changes).
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(doc_bytes))
    changes = []

    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue
        first_row = table.rows[0]
        tr = first_row._tr
        trPr = tr.get_or_add_trPr()

        # Check if tblHeader already set
        if trPr.find(qn("w:tblHeader")) is not None:
            continue

        header_elem = OxmlElement("w:tblHeader")
        trPr.append(header_elem)
        changes.append(f"Table {i+1}: marked first row as header")

    if not changes:
        return doc_bytes, changes

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), changes


# ── Fix: document language ────────────────────────────────────────────────────

def fix_language(doc_bytes: bytes, lang: str = "en-US") -> tuple[bytes, list[str]]:
    """
    Set the document's default paragraph language to lang (e.g. 'en-US').
    Returns (updated_docx_bytes, list_of_changes).
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(doc_bytes))
    changes = []
    set_count = 0

    # Apply to every paragraph run that has no language set
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._r.get_or_add_rPr()
            lang_elem = rPr.find(qn("w:lang"))
            if lang_elem is None:
                lang_elem = OxmlElement("w:lang")
                rPr.append(lang_elem)
            if not lang_elem.get(qn("w:val")):
                lang_elem.set(qn("w:val"), lang)
                set_count += 1

    if set_count:
        changes.append(f"Set language '{lang}' on {set_count} text run(s)")

    # Also set in document styles default
    try:
        from docx.enum.style import WD_STYLE_TYPE
        normal_style = next(
            (s for s in doc.styles if s.name == "Normal"
             and s.type == WD_STYLE_TYPE.PARAGRAPH), None
        )
        if normal_style is None:
            raise LookupError("Normal style not found")
        rPr = normal_style.element.get_or_add_rPr()
        lang_elem = rPr.find(qn("w:lang"))
        if lang_elem is None:
            lang_elem = OxmlElement("w:lang")
            rPr.append(lang_elem)
        if lang_elem.get(qn("w:val")) != lang:
            lang_elem.set(qn("w:val"), lang)
            changes.append(f"Set language '{lang}' on Normal style")
    except Exception:
        pass

    if not changes:
        return doc_bytes, changes

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), changes


# ── Fix: heading order ────────────────────────────────────────────────────────

def fix_heading_order(doc_bytes: bytes) -> tuple[bytes, list[str]]:
    """
    Normalize heading level sequence so no level is skipped.
    Returns (updated_docx_bytes, list_of_changes).
    """
    from docx import Document

    doc = Document(io.BytesIO(doc_bytes))
    changes = []
    prev_level = 0

    for para in doc.paragraphs:
        style = para.style.name  # e.g. 'Heading 1', 'Heading 2'
        if not style.startswith("Heading "):
            continue
        try:
            level = int(style.split()[-1])
        except ValueError:
            continue

        if prev_level > 0 and level > prev_level + 1:
            new_level = prev_level + 1
            new_style = f"Heading {new_level}"
            try:
                para.style = doc.styles[new_style]
                changes.append(f"Renormalized '{style}' → '{new_style}'")
                level = new_level
            except Exception as e:
                changes.append(f"Could not apply {new_style}: {e}")
        prev_level = level

    if not changes:
        return doc_bytes, changes

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), changes


# ── Fix: headings presence ────────────────────────────────────────────────────

def fix_headings_presence(doc_bytes: bytes) -> tuple[bytes, list[str]]:
    """
    If the document has no Heading-styled paragraphs, restyle the first
    non-empty paragraph as Heading 1.
    Returns (updated_docx_bytes, list_of_changes).
    """
    from docx import Document

    doc = Document(io.BytesIO(doc_bytes))
    has_heading = any(
        p.style.name.startswith("Heading") for p in doc.paragraphs
    )
    if has_heading:
        return doc_bytes, []

    for para in doc.paragraphs:
        if para.text.strip():
            try:
                para.style = doc.styles["Heading 1"]
            except KeyError:
                # Heading 1 style missing — create a minimal one with outline level 0
                # so screen readers and Ally recognise it as H1.
                from docx.enum.style import WD_STYLE_TYPE
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                try:
                    h1 = doc.styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
                    pPr = h1.element.get_or_add_pPr()
                    outline = OxmlElement("w:outlineLvl")
                    outline.set(qn("w:val"), "0")
                    pPr.append(outline)
                    para.style = h1
                except Exception as e:
                    print(f"    WARNING: Could not create Heading 1 style: {e}")
                    return doc_bytes, []
            label = para.text[:60]
            changes = [
                f"Restyled first paragraph as Heading 1 (placeholder — please review): '{label}'"
            ]
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue(), changes

    return doc_bytes, []  # empty document


# ── Fix: headings start at one ────────────────────────────────────────────────

def fix_headings_start_at_one(doc_bytes: bytes) -> tuple[bytes, list[str]]:
    """
    If the document has no H1 heading, insert a placeholder H1 paragraph at
    the very beginning of the document.  This preserves the existing heading
    hierarchy (e.g. H2/H3 stay as-is) rather than shifting every level, which
    could scramble structured documents like exams or reports.
    Returns (updated_docx_bytes, list_of_changes).
    """
    import re
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document(io.BytesIO(doc_bytes))

    heading_paras = [
        (p, int(re.search(r"\d+", p.style.name).group()))
        for p in doc.paragraphs
        if re.match(r"Heading \d+", p.style.name)
    ]
    if not heading_paras:
        return doc_bytes, []

    min_level = min(level for _, level in heading_paras)

    # Find the first non-empty paragraph
    first_nonempty = next((p for p in doc.paragraphs if p.text.strip()), None)
    if first_nonempty is None:
        return doc_bytes, []

    first_is_heading1 = first_nonempty.style.name == "Heading 1"

    # Nothing to do: document already opens with H1
    if min_level == 1 and first_is_heading1:
        return doc_bytes, []

    # Idempotency: placeholder already inserted from a prior run
    if (first_nonempty.style.name == "Heading 1" and
            "placeholder" in first_nonempty.text.lower()):
        return doc_bytes, []

    # Get or create Heading 1 style
    try:
        doc.styles["Heading 1"]
    except KeyError:
        try:
            h1_style = doc.styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
            pPr = h1_style.element.get_or_add_pPr()
            outline = OxmlElement("w:outlineLvl")
            outline.set(qn("w:val"), "0")
            pPr.append(outline)
        except Exception as e:
            return doc_bytes, [f"Could not create Heading 1 style: {e}"]

    placeholder_text = "Document Title (placeholder — please replace with actual title)"

    if min_level > 1:
        # No H1 exists at all — insert placeholder before the first heading
        reason = f"document started at H{min_level}"
        target = heading_paras[0][0]
    else:
        # H1 exists but is not the first paragraph — insert placeholder at top
        reason = "first paragraph is not Heading 1"
        target = first_nonempty

    new_para_xml = OxmlElement("w:p")
    new_pPr = OxmlElement("w:pPr")
    new_pStyle = OxmlElement("w:pStyle")
    new_pStyle.set(qn("w:val"), "Heading1")
    new_pPr.append(new_pStyle)
    new_para_xml.append(new_pPr)
    new_r = OxmlElement("w:r")
    new_t = OxmlElement("w:t")
    new_t.text = placeholder_text
    new_r.append(new_t)
    new_para_xml.append(new_r)
    target._element.addprevious(new_para_xml)

    changes = [f"Inserted placeholder Heading 1 at top ({reason})"]
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), changes


# ── Orchestrate fixes on one file ─────────────────────────────────────────────

def fix_word_file(client: CanvasClient, course_id: int, file_info: dict,
                  fixes: list[str], dry_run: bool = False) -> dict:
    """
    Download a Word file, apply fixes, re-upload to Canvas.
    Returns result dict.
    """
    file_id = file_info["id"]
    filename = file_info.get("display_name") or file_info.get("filename", "document.docx")
    folder_id = file_info.get("folder_id")

    print(f"  Processing: {filename}")
    doc_bytes = client.download_url(file_info["url"])
    all_changes = []
    all_fixes = "all" in fixes

    if all_fixes or "headings_presence" in fixes:
        doc_bytes, ch = fix_headings_presence(doc_bytes)
        all_changes.extend(ch)

    if all_fixes or "headings_start_at_one" in fixes:
        doc_bytes, ch = fix_headings_start_at_one(doc_bytes)
        all_changes.extend(ch)

    if all_fixes or "heading_order" in fixes:
        doc_bytes, ch = fix_heading_order(doc_bytes)
        all_changes.extend(ch)

    if all_fixes or "table_headers" in fixes:
        doc_bytes, ch = fix_table_headers(doc_bytes)
        all_changes.extend(ch)

    if all_fixes or "no_language" in fixes or "language" in fixes:
        doc_bytes, ch = fix_language(doc_bytes)
        all_changes.extend(ch)

    if all_fixes or "image_alt" in fixes:
        doc_bytes, ch = fix_image_alt(doc_bytes, filename, placeholder=False)
        all_changes.extend(ch)

    if "image_alt_placeholder" in fixes:
        doc_bytes, ch = fix_image_alt(doc_bytes, filename, placeholder=True)
        all_changes.extend(ch)

    result = {"file": filename, "file_id": file_id, "changes": all_changes, "updated": False}

    if all_changes and not dry_run:
        client.upload_file(course_id, folder_id, filename, WORD_MIME, doc_bytes)
        result["updated"] = True
        print(f"    Re-uploaded with {len(all_changes)} fix(es)")
    elif all_changes:
        print(f"    [dry-run] {len(all_changes)} change(s) would be made")
    else:
        print(f"    No changes needed")

    for ch in all_changes:
        print(f"      • {ch}")

    return result


def fix_course_word_files(client: CanvasClient, course_id: int,
                          fixes: list[str], file_id: int = None,
                          dry_run: bool = False) -> list[dict]:
    """Run fixes on one Word file or all Word files in a course."""
    if file_id:
        files = [client.get_file_info(file_id)]
    else:
        all_files = client.list_files(course_id)
        files = [
            f for f in all_files
            if Path(f.get("filename", "")).suffix.lower() in WORD_EXTS
        ]
        print(f"Found {len(files)} Word file(s) in course {course_id}")

    results = []
    for f in files:
        try:
            r = fix_word_file(client, course_id, f, fixes, dry_run)
            results.append(r)
            time.sleep(0.5)
        except Exception as e:
            print(f"  ERROR on '{f.get('filename', '?')}': {e}")
            results.append({"file": f.get("filename", "?"), "error": str(e)})

    return results


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Fix Word document accessibility issues")
    parser.add_argument("--course-id", type=int, required=True)
    parser.add_argument("--file-id", type=int, default=None,
                        help="Canvas file ID of a specific .docx, or omit for all")
    parser.add_argument("--fix", default="all",
                        help="Comma-separated: all,image_alt,image_alt_placeholder,"
                             "table_headers,no_language,heading_order,"
                             "headings_presence,headings_start_at_one")
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    fix_list = [f.strip() for f in args.fix.split(",")]
    client = CanvasClient()

    results = fix_course_word_files(client, args.course_id, fix_list,
                                    file_id=args.file_id, dry_run=args.dry_run)

    total_changes = sum(len(r.get("changes", [])) for r in results)
    updated = sum(1 for r in results if r.get("updated"))
    print(f"\nSummary: {total_changes} changes, {updated}/{len(results)} files updated")


if __name__ == "__main__":
    main()
